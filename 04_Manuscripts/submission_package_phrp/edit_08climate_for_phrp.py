"""
NDB_XXX_temperature_CVD – PHRP投稿版 DOCX 生成スクリプト
================================================================
08Climate_...docx を直接編集して PHRP 形式の Manuscript_PHRP.docx を生成する。

修正内容
--------
[事実誤り]
  (A) Abstract Methods: "fiscal year 2020" → "fiscal year 2023"
  (B) Methods §1:       "fiscal year 2020 (April 2020 through March 2021)"
                      → "fiscal year 2023 (April 2023 through March 2024)"
  (C) Methods §2:       "released 2023" → "released 2025"
                        ※ 第10回NDBオープンデータは2025年5月公開

[PHRP 投稿規定]
  (D) Abstract Objective 見出し: "Objective" → "Objectives"
  (E) Abstract Conclusions 見出し: "Conclusions" → "Conclusion"
  (F) Keywords: 7 語 → 6 語・アルファベット順
  (G) AI 開示文: Methods/Ethics Statement の後に追加

[ファイル操作]
  - 既存 Manuscript_PHRP.docx をアーカイブ（日付付きリネーム）
  - 修正済みを Manuscript_PHRP.docx として保存
"""

import os
import shutil
from datetime import datetime
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------
# パス設定
# ---------------------------------------------------------------
BASE = os.path.dirname(os.path.abspath(__file__))           # submission_package_phrp/
MANUSCRIPTS_DIR = os.path.dirname(BASE)                      # 04_Manuscripts/

SOURCE_DOCX = os.path.join(
    MANUSCRIPTS_DIR,
    "08Climate_Sensitive_Cardiovascular_Pharmaceutical_Demand_final_refs_figure_checked.docx",
)
OUTPUT_DOCX = os.path.join(BASE, "Manuscript_PHRP.docx")
ARCHIVE_DOCX = os.path.join(
    BASE,
    f"Manuscript_PHRP_v1_archived_{datetime.now().strftime('%Y%m%d')}.docx",
)

# ---------------------------------------------------------------
# Step 0: 既存 Manuscript_PHRP.docx をアーカイブ
# ---------------------------------------------------------------
if os.path.exists(OUTPUT_DOCX):
    shutil.copy2(OUTPUT_DOCX, ARCHIVE_DOCX)
    print(f"[ARCHIVE] {os.path.basename(ARCHIVE_DOCX)}")
else:
    print("[INFO] 既存 Manuscript_PHRP.docx なし（アーカイブ不要）")

# ---------------------------------------------------------------
# Step 1: ソース DOCX を読み込む
# ---------------------------------------------------------------
print(f"\n[LOAD] {os.path.basename(SOURCE_DOCX)}")
doc = Document(SOURCE_DOCX)

# ---------------------------------------------------------------
# ヘルパー: パラグラフ内の runs で old→new 置換
# ---------------------------------------------------------------
def replace_in_para(para, old: str, new: str) -> bool:
    """
    runs を走査して old→new を置換する。
    単一 run 内にある場合はその run のみ変更。
    複数 run にまたがる場合は全テキストを先頭 run に集約（書式はその run の書式を引き継ぐ）。
    """
    full_text = "".join(r.text for r in para.runs)
    if old not in full_text:
        return False

    # --- シングル run 内に完全に存在する場合 ---
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True

    # --- マルチ run にまたがる場合: 先頭 run に集約 ---
    new_full = full_text.replace(old, new)
    if not para.runs:
        return False
    para.runs[0].text = new_full
    for r in para.runs[1:]:
        r.text = ""
    return True


def apply_corrections(paragraphs, corrections):
    """パラグラフリストに対して corrections を適用し、変更件数を返す。"""
    count = 0
    for para in paragraphs:
        for old, new in corrections:
            if replace_in_para(para, old, new):
                print(f"  [FIX] '{old[:70]}' → '{new[:70]}'")
                count += 1
    return count


# ---------------------------------------------------------------
# Step 2: 事実誤り修正 (A)(B)(C)
# ---------------------------------------------------------------
print("\n--- 事実誤り修正 ---")

FACTUAL_CORRECTIONS = [
    # (A) Abstract Methods – short phrase
    (
        "NDB Open Data for fiscal year 2020",
        "NDB Open Data for fiscal year 2023",
    ),
    # (B) Methods §1 – long phrase（先に長いものを処理）
    (
        "fiscal year 2020 (April 2020 through March 2021)",
        "fiscal year 2023 (April 2023 through March 2024)",
    ),
    # 念のため残った "fiscal year 2020" を全部変換
    (
        "fiscal year 2020",
        "fiscal year 2023",
    ),
    # (C) Methods §2  ※第10回 NDB は 2025 年 5 月公開
    (
        "released 2023",
        "released 2025",
    ),
]

n = apply_corrections(doc.paragraphs, FACTUAL_CORRECTIONS)
# テーブル内も修正
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            n += apply_corrections(cell.paragraphs, FACTUAL_CORRECTIONS)
print(f"  合計 {n} 箇所を修正")

# ---------------------------------------------------------------
# Step 3: PHRP 書式修正
# ---------------------------------------------------------------
print("\n--- PHRP 書式修正 ---")

# (D)(E) Abstract 内の見出し変更
in_abstract = False
abstract_conclusions_fixed = False

for para in doc.paragraphs:
    if para.style.name == "Heading 1" and para.text.strip() == "Abstract":
        in_abstract = True
        continue
    if para.style.name == "Heading 1" and para.text.strip() != "Abstract":
        in_abstract = False

    if in_abstract and para.style.name == "Heading 2":
        text = para.text.strip()

        # (D) Objective → Objectives
        if text == "Objective":
            replace_in_para(para, "Objective", "Objectives")
            print("  [FIX] Abstract見出し: 'Objective' → 'Objectives'")

        # (E) Conclusions → Conclusion（Abstract 内のみ・最初の1回）
        if text == "Conclusions" and not abstract_conclusions_fixed:
            replace_in_para(para, "Conclusions", "Conclusion")
            print("  [FIX] Abstract見出し: 'Conclusions' → 'Conclusion'")
            abstract_conclusions_fixed = True

# (F) Keywords: 7 語 → 6 語、アルファベット順
# 現在: "climate; cardiovascular disease; pharmaceutical demand; prescription rate;
#        ecological study; NDB Open Data; Japan"  (7 語)
# 修正: "cardiovascular disease; climate; ecological study; Japan; NDB Open Data;
#        prescription rate"                         (6 語、アルファベット順)
NEW_KEYWORDS = (
    "cardiovascular disease; climate; ecological study; "
    "Japan; NDB Open Data; prescription rate"
)
for para in doc.paragraphs:
    if "Keywords:" in para.text:
        # "Keywords:" ラベル (Bold) の次の run を更新
        keyword_found = False
        for run in para.runs:
            if "Keywords:" in run.text:
                keyword_found = True
                continue
            if keyword_found and run.text.strip():
                run.text = NEW_KEYWORDS
                print(f"  [FIX] Keywords: 7 語 → 6 語（アルファベット順）")
                # 残りの run をクリア
                in_kw = False
                for r2 in para.runs:
                    if r2 is run:
                        in_kw = True
                        continue
                    if in_kw:
                        r2.text = ""
                break
        break

# (G) AI 開示文: Methods § 5.6 Ethics Statement の後に追加
# スタイル名 "Heading3" か "Heading 3" は環境によって異なるため両方チェック
inserted = False
for i, para in enumerate(doc.paragraphs):
    if (
        para.style.name in ("Heading 3", "Heading3")
        and "Ethics Statement" in para.text
    ):
        # Ethics Statement 本文の次の段落の後ろに挿入
        ethics_body = doc.paragraphs[i + 1] if i + 1 < len(doc.paragraphs) else para

        # --- 新しい見出し § 5.7 ---
        new_heading = OxmlElement("w:p")
        new_hPr = OxmlElement("w:pPr")
        new_hStyle = OxmlElement("w:pStyle")
        new_hStyle.set(qn("w:val"), para.style.name.replace(" ", ""))  # "Heading3"
        new_hPr.append(new_hStyle)
        new_heading.append(new_hPr)
        new_h_r = OxmlElement("w:r")
        new_h_t = OxmlElement("w:t")
        new_h_t.text = "5.7. AI Disclosure"
        new_h_r.append(new_h_t)
        new_heading.append(new_h_r)

        # --- 新しい本文 ---
        new_body = OxmlElement("w:p")
        new_bPr = OxmlElement("w:pPr")
        new_bStyle = OxmlElement("w:pStyle")
        # 本文スタイルは "FirstParagraph" か "First Paragraph" か確認
        body_style = doc.paragraphs[i + 1].style.name if i + 1 < len(doc.paragraphs) else "Body Text"
        new_bStyle.set(qn("w:val"), body_style.replace(" ", ""))
        new_bPr.append(new_bStyle)
        new_body.append(new_bPr)
        new_b_r = OxmlElement("w:r")
        new_b_t = OxmlElement("w:t")
        new_b_t.set(qn("xml:space"), "preserve")
        new_b_t.text = (
            "Artificial intelligence (AI) tools (Claude, Anthropic) were used to assist "
            "with manuscript drafting, code generation, and language editing. "
            "All analytical decisions, interpretation of results, and final manuscript "
            "content were reviewed and approved by the authors."
        )
        new_b_r.append(new_b_t)
        new_body.append(new_b_r)

        # Ethics 本文の後に挿入（heading → body の順）
        ethics_body._element.addnext(new_body)
        ethics_body._element.addnext(new_heading)

        print("  [ADD] § 5.7 AI Disclosure を Ethics Statement の後に挿入")
        inserted = True
        break

if not inserted:
    print("  [WARN] Ethics Statement 見出しが見つからなかった。AI Disclosure 未追加。")

# ---------------------------------------------------------------
# Step 4: 保存
# ---------------------------------------------------------------
doc.save(OUTPUT_DOCX)
print(f"\n[SAVED] {os.path.basename(OUTPUT_DOCX)}")

# ---------------------------------------------------------------
# Step 5: 確認サマリ
# ---------------------------------------------------------------
print("\n========================================")
print("修正サマリ")
print("========================================")
print("[事実誤り修正]")
print("  (A) Abstract Methods: 'fiscal year 2020' → 'fiscal year 2023'")
print("  (B) Methods §1: '...April 2020 through March 2021...'")
print("            → '...April 2023 through March 2024...'")
print("  (C) Methods §2: 'released 2023' → 'released 2025'")
print("      ※ 第10回NDBオープンデータ = 2025年5月公開")
print("[PHRP 書式修正]")
print("  (D) Abstract見出し: 'Objective' → 'Objectives'")
print("  (E) Abstract見出し: 'Conclusions' → 'Conclusion'")
print("  (F) Keywords: 7語 → 6語（アルファベット順）")
print("       cardiovascular disease; climate; ecological study;")
print("       Japan; NDB Open Data; prescription rate")
print("  (G) § 5.7 AI Disclosure 追加")
print("\n[次に対応が必要なファイル]")
print("  - Figure1_choropleth_cvd_percapita.png  : タイトル内 '(FY2020)' → '(FY2023)'")
print("  - TitlePage_PHRP.docx                  : FY記載の確認・更新")
print("  - config/config.yaml                   : ndb_year: 2020 → 2023、コメント修正")
print("  - README.md / README_ja.md             : FY2020 → FY2023、GitHub/Zenodo設定")
print("  - GitHub_Zenodo_setup.md               : FY2020 → FY2023")
