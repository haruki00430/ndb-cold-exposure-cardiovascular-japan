"""
STROBE チェックリスト DOCX 生成スクリプト
観察的研究（生態学的研究）用 STROBE Statement v4
"""

import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(SCRIPT_DIR, "STROBE_checklist_PHRP.docx")


def make_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10)
    return doc


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11) if level == 1 else Pt(10)
    run.bold = True
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def para(doc, text):
    p = doc.add_paragraph(text)
    if p.runs:
        p.runs[0].font.name = 'Times New Roman'
        p.runs[0].font.size = Pt(10)
    return p


def add_row(table, item, recommendation, location):
    row = table.add_row()
    for i, txt in enumerate([item, recommendation, location]):
        cell = row.cells[i]
        cell.text = txt
        for r in cell.paragraphs[0].runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(9)
    return row


def build():
    doc = make_doc()

    heading(doc, "STROBE Statement — Checklist of Items That Should Be Included in Reports "
                 "of Observational Studies (Ecological Study)")
    para(doc, "")
    para(doc, "Manuscript: Association between Cold Month Exposure and Cardiovascular Disease "
              "Prescription Rate: A Nationwide Ecological Study in Japan")
    para(doc, "Journal: Osong Public Health and Research Perspectives (PHRP)")
    para(doc, "")

    # テーブル作成
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    # ヘッダー
    hdr = table.rows[0].cells
    for i, h in enumerate(["Item No & Recommendation", "Checklist Item", "Reported on page No."]):
        hdr[i].text = h
        for r in hdr[i].paragraphs[0].runs:
            r.bold = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(9)

    # チェックリスト項目
    items = [
        # TITLE AND ABSTRACT
        ("TITLE AND ABSTRACT", "", ""),
        ("1 (a)", "Indicate the study's design with a commonly used term in the title or the abstract",
         "Title, Abstract (Methods)"),
        ("1 (b)", "Provide in the abstract an informative and balanced summary of what was done and what was found",
         "Abstract (pp. 1)"),

        # INTRODUCTION
        ("INTRODUCTION", "", ""),
        ("2", "Background/rationale: Explain the scientific background and rationale for the investigation being reported",
         "Introduction (pp. 2–3)"),
        ("3", "Objectives: State specific objectives, including any prespecified hypotheses",
         "Introduction (last paragraph, p. 3)"),

        # METHODS
        ("METHODS", "", ""),
        ("4", "Study design: Present key elements of study design early in the paper",
         "Methods §1 (p. 3)"),
        ("5", "Setting: Describe the setting, locations, and relevant dates, including periods of recruitment, exposure, follow-up, and data collection",
         "Methods §1–2 (p. 3)"),
        ("6 (a)", "Participants (cohort): Give the eligibility criteria, and the sources and methods of selection of participants",
         "Methods §1 (p. 3); N = 47 prefectures"),
        ("6 (b)", "For matched studies, give matching criteria and number of exposed and unexposed",
         "N/A"),
        ("7", "Variables: Clearly define all outcomes, exposures, predictors, potential confounders, and effect modifiers. Give diagnostic criteria, if applicable",
         "Methods §2–4 (pp. 3–4)"),
        ("8", "Data sources/measurement: For each variable of interest, give sources of data and details of methods of assessment (measurement). Describe comparability of assessment methods if there is more than one group",
         "Methods §2–4 (pp. 3–4)"),
        ("9", "Bias: Describe any efforts to address potential sources of bias",
         "Discussion §8 / Limitations (p. 8)"),
        ("10", "Study size: Explain how the study size was arrived at",
         "Methods §1 (N = 47, all prefectures)"),
        ("11", "Quantitative variables: Explain how quantitative variables were handled in the analyses. If applicable, describe which groupings were chosen and why",
         "Methods §5 (pp. 4–5)"),
        ("12 (a)", "Statistical methods: Describe all statistical methods, including those used to control for confounding",
         "Methods §5.1–5.4 (pp. 4–5)"),
        ("12 (b)", "Describe any methods used to examine subgroups and interactions",
         "Methods §5.2 (model selection)"),
        ("12 (c)", "Explain how missing data were addressed",
         "Methods §2 (NDB: complete prefectural data)"),
        ("12 (d)", "If applicable, explain how loss to follow-up was addressed",
         "N/A (ecological study, cross-sectional)"),
        ("12 (e)", "Describe any sensitivity analyses",
         "Not applicable (pre-specified primary exposure)"),

        # RESULTS
        ("RESULTS", "", ""),
        ("13 (a)", "Participants: Report numbers of individuals at each stage of study — e.g. numbers potentially eligible, examined for eligibility, confirmed eligible, included in the study, completing follow-up, and analysed",
         "Methods §1; Results §1 (N = 47)"),
        ("13 (b)", "Give reasons for non-participation at each stage",
         "N/A (all 47 prefectures included)"),
        ("13 (c)", "Consider use of a flow diagram",
         "Not included (all prefectures eligible)"),
        ("14 (a)", "Descriptive data: Give characteristics of study participants (e.g. demographic, clinical, social) and information on exposures and potential confounders",
         "Results §1 / Table 1 (p. 5)"),
        ("14 (b)", "Indicate number of participants with missing data for each variable of interest",
         "Results §1 (no missing data for main variables)"),
        ("14 (c)", "Summarise follow-up time (e.g. average and total amount)",
         "N/A (cross-sectional)"),
        ("15", "Outcome data: Report numbers of outcome events or summary measures over time",
         "Results §1 / Table 1 (p. 5)"),
        ("16 (a)", "Main results: Give unadjusted estimates and, if applicable, confounder-adjusted estimates and their precision",
         "Results §3–4 / Tables 3–4 (pp. 5–6)"),
        ("16 (b)", "Report category boundaries when continuous variables were categorized",
         "Methods §3 (10°C threshold)"),
        ("16 (c)", "If relevant, consider translating estimates of relative risk into absolute risk for a meaningful time period",
         "Results §4 (β coefficients reported)"),
        ("17", "Other analyses: Report other analyses done — e.g. analyses of subgroups and interactions, and sensitivity analyses",
         "Results §5 / Table 5 (spatial autocorrelation)"),

        # DISCUSSION
        ("DISCUSSION", "", ""),
        ("18", "Key results: Summarise key results with reference to study objectives",
         "Discussion §1 (p. 6)"),
        ("19", "Limitations: Discuss limitations of the study, taking into account sources of potential bias or imprecision. Discuss both direction and magnitude of any potential bias",
         "Discussion §8 (p. 8)"),
        ("20", "Interpretation: Give a cautious overall interpretation of results considering objectives, limitations, multiplicity of analyses, results from similar studies, and other relevant evidence",
         "Discussion §2–6 (pp. 6–8)"),
        ("21", "Generalisability: Discuss the generalisability (external validity) of the study results",
         "Discussion §7–8 (pp. 7–8)"),

        # OTHER INFORMATION
        ("OTHER INFORMATION", "", ""),
        ("22", "Funding: Give the source of funding and the role of the funders for the present study and, if applicable, for the original study on which the present article is based",
         "Title Page (Notes: Funding)"),
    ]

    for item_no, checklist_item, location in items:
        if checklist_item == "" and location == "":
            # セクション見出し行
            row = table.add_row()
            cell = row.cells[0]
            cell.merge(row.cells[1])
            cell.merge(row.cells[2])
            cell.text = item_no
            for r in cell.paragraphs[0].runs:
                r.bold = True
                r.font.name = 'Times New Roman'
                r.font.size = Pt(9)
                from docx.shared import RGBColor
                r.font.color.rgb = RGBColor(0x2F, 0x5F, 0x98)
        else:
            add_row(table, item_no, checklist_item, location)

    # 脚注
    para(doc, "")
    p = doc.add_paragraph()
    p.add_run(
        "Note: ").bold = True
    p.add_run(
        "This checklist is based on the STROBE Statement for observational studies "
        "(von Elm E, et al. Lancet 2007;370:1453–1457). "
        "For ecological studies, item 6 refers to the unit of analysis (prefecture) rather than individual participants. "
        "N/A = not applicable.")
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(9)

    doc.save(OUT_PATH)
    print(f"[OK] {OUT_PATH}")


if __name__ == "__main__":
    build()
