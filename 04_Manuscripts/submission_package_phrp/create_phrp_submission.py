"""
PHRP (Osong Public Health and Research Perspectives) 投稿パッケージ生成スクリプト
- Manuscript_PHRP.docx  : ブラインド原稿（著者情報なし）
- TitlePage_PHRP.docx   : タイトルページ（著者情報・Notes含む）
- CoverLetter_PHRP.docx : カバーレター
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

# ==============================================================================
# パス設定
# ==============================================================================
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_DIR = SCRIPT_DIR  # submission_package_phrp/ に出力


def set_phrp_style(doc):
    """PHRP基本スタイル: Times New Roman 11pt, A4, 2.5cm margins, double-space"""
    from docx.shared import Inches
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # デフォルトスタイル
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)


def add_double_spaced_paragraph(doc, text='', bold=False, center=False, heading=False):
    """ダブルスペース段落を追加"""
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    para = doc.add_paragraph()
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif not heading:
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ダブルスペース設定
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:line'), '480')  # 240 = single, 480 = double
    spacing.set(qn('w:lineRule'), 'auto')
    spacing.set(qn('w:after'), '0')
    pPr.append(spacing)

    if text:
        run = para.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.bold = bold

    return para


def add_run_to_para(para, text, bold=False, italic=False):
    """段落にrunを追加"""
    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    return run


def add_heading_para(doc, text, level=1):
    """見出し段落（ボールド、センタリング）"""
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:line'), '480')
    spacing.set(qn('w:lineRule'), 'auto')
    spacing.set(qn('w:after'), '0')
    pPr.append(spacing)

    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.bold = True
    return para


def add_section_heading(doc, text):
    """セクション見出し（ボールド、センタリング）"""
    return add_heading_para(doc, text)


def add_subsection_heading(doc, text):
    """サブセクション見出し（ボールド、左寄せ）"""
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:line'), '480')
    spacing.set(qn('w:lineRule'), 'auto')
    spacing.set(qn('w:after'), '0')
    pPr.append(spacing)

    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.bold = True
    return para


def add_page_break(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(WD_ALIGN_PARAGRAPH.LEFT.__class__.__mro__[0])
    from docx.oxml.ns import qn
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)


def add_table(doc, headers, rows, caption=None):
    """APA-styleテーブルを追加"""
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    if caption:
        p = add_double_spaced_paragraph(doc, caption, bold=True)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # ヘッダー行
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)

    # データ行
    for ri, row_data in enumerate(rows):
        row_cells = table.rows[ri + 1].cells
        for ci, cell_text in enumerate(row_data):
            row_cells[ci].text = str(cell_text)
            for run in row_cells[ci].paragraphs[0].runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)

    return table


# ==============================================================================
# Manuscript_PHRP.docx（ブラインド原稿）
# ==============================================================================
def create_manuscript(out_path):
    doc = Document()
    set_phrp_style(doc)

    # ---- タイトル ----
    p = add_double_spaced_paragraph(doc, bold=True, center=True)
    add_run_to_para(p,
        "Association between Cold Month Exposure and Cardiovascular Disease Prescription Rate: "
        "A Nationwide Ecological Study in Japan",
        bold=True)

    add_double_spaced_paragraph(doc)  # blank line

    # ---- Abstract ----
    add_section_heading(doc, "Abstract")
    add_double_spaced_paragraph(doc)

    # Objectives
    p = add_double_spaced_paragraph(doc)
    add_run_to_para(p, "Objectives: ", bold=True)
    add_run_to_para(p,
        "Cardiovascular diseases (CVDs) impose a substantial burden in Japan, and ambient temperature "
        "is a recognized environmental risk factor. We examined whether cold month exposure was "
        "associated with prefectural-level CVD prescription rates (per 100,000 population) across Japan, "
        "independent of socioeconomic confounders.")

    # Methods
    p = add_double_spaced_paragraph(doc)
    add_run_to_para(p, "Methods: ", bold=True)
    add_run_to_para(p,
        "We conducted a nationwide ecological study using 47 Japanese prefectures as the unit of analysis. "
        "CVD prescription rates were derived from the 10th National Database (NDB) Open Data (fiscal year 2023) "
        "and standardized by 2020 National Census population. Cold month ratio (proportion of months with mean "
        "temperature <10°C) was calculated from Japan Meteorological Agency data. Covariates included "
        "physicians per 100,000, unemployment rate, aging rate, and per capita prefectural income. "
        "Multivariable ordinary least squares (OLS) regression was used. Spatial autocorrelation was "
        "assessed using Global Moran’s I.")

    # Results
    p = add_double_spaced_paragraph(doc)
    add_run_to_para(p, "Results: ", bold=True)
    add_run_to_para(p,
        "The mean CVD prescription rate across prefectures was 17,521,396 ± 2,607,830 per 100,000 "
        "population. Cold month ratio showed the strongest univariable association with CVD prescription "
        "rate (R² = 0.390, ")
    add_run_to_para(p, "p", italic=True)
    add_run_to_para(p, " < 0.001) among five temperature indicators. In the multivariable model "
        "(R² = 0.667, adjusted R² = 0.626), cold month ratio (β = +10,450,255; ")
    add_run_to_para(p, "p", italic=True)
    add_run_to_para(p, " < 0.001), aging rate (β = +499,769; ")
    add_run_to_para(p, "p", italic=True)
    add_run_to_para(p, " < 0.001), and unemployment rate (β = +1,300,753; ")
    add_run_to_para(p, "p", italic=True)
    add_run_to_para(p, " = 0.025) were independently and significantly associated with higher CVD "
        "prescription rates. Residual diagnostics confirmed model adequacy. Global Moran’s I was "
        "not significant (I = 0.091, ")
    add_run_to_para(p, "p", italic=True)
    add_run_to_para(p, " = 0.114).")

    # Conclusion
    p = add_double_spaced_paragraph(doc)
    add_run_to_para(p, "Conclusion: ", bold=True)
    add_run_to_para(p,
        "Cold month exposure was independently and strongly associated with higher CVD prescription rates "
        "across Japanese prefectures. Prefectures with prolonged cold seasons bear a disproportionate CVD "
        "pharmaceutical burden, warranting targeted seasonal prevention strategies and healthcare resource "
        "allocation.")

    add_double_spaced_paragraph(doc)

    # Keywords
    p = add_double_spaced_paragraph(doc)
    add_run_to_para(p, "Keywords: ", bold=True)
    add_run_to_para(p, "cardiovascular disease; cold exposure; ecological study; Japan; "
                       "NDB Open Data; prescription rate")

    add_double_spaced_paragraph(doc)

    # ---- Introduction ----
    add_section_heading(doc, "Introduction")
    add_double_spaced_paragraph(doc)

    intro_paras = [
        ("Cardiovascular diseases (CVDs), including ischemic heart disease and cerebrovascular disease, "
         "remain the leading cause of mortality worldwide, accounting for approximately 19.8 million deaths "
         "in 2022 [1]. In Japan, CVDs represent the second leading cause of death after malignant neoplasms "
         "and impose a substantial economic burden on the healthcare system [2]. Understanding environmental "
         "determinants of CVD burden at the population level is essential for informing public health "
         "prevention strategies and healthcare resource allocation."),

        ("Ambient temperature is a well-established risk factor for cardiovascular morbidity and mortality. "
         "Epidemiological studies have consistently demonstrated that cold exposure increases the risk of "
         "acute cardiovascular events, including myocardial infarction, stroke, and heart failure "
         "exacerbation [3–5]. The physiological mechanisms are well characterized: cold exposure "
         "increases sympathetic nervous system activity, elevates blood pressure, increases blood viscosity, "
         "and promotes thrombogenesis [6,7]. A large multi-country study by Gasparrini et al. demonstrated "
         "that cold-related mortality substantially exceeded heat-related mortality in most temperate "
         "countries [8]. Globally, Zhao et al. estimated that non-optimal ambient temperatures accounted "
         "for 5.08 million excess deaths annually during 2000–2019, with cold-related mortality "
         "predominating in temperate regions including East Asia [9]."),

        ("However, the majority of studies on temperature and CVD have focused on acute clinical outcomes "
         "(emergency department visits, hospitalizations, mortality) at the individual level. The "
         "relationship between cumulative cold exposure and CVD-related drug prescription patterns—which "
         "reflects chronic disease management and ongoing healthcare utilization—has received limited "
         "attention. Drug prescription rates serve as a proxy indicator of CVD prevalence and treatment "
         "intensity at the population level, capturing not only acute disease burden but also the "
         "pharmacological management of chronic cardiovascular conditions such as hypertension, dyslipidemia, "
         "and heart failure."),

        ("Japan’s archipelago spans approximately 25 degrees of latitude, from subarctic Hokkaido to "
         "subtropical Okinawa, with the proportion of cold months (mean monthly temperature <10°C) "
         "varying substantially across prefectures [10]. This geographic gradient in cold exposure, combined "
         "with the availability of the National Database (NDB) Open Data—a comprehensive administrative "
         "database covering virtually all health insurance claims in Japan—offers a unique opportunity "
         "to examine the cold exposure–CVD prescription relationship at a nationwide ecological scale."),

        ("In this study, we examined whether cold month ratio was associated with prefectural-level CVD "
         "prescription rates (per 100,000 population) across Japan, independent of socioeconomic covariates, "
         "using multivariable regression. We also assessed spatial autocorrelation to determine whether CVD "
         "prescription patterns exhibited geographic clustering."),
    ]
    for txt in intro_paras:
        add_double_spaced_paragraph(doc, txt)

    add_double_spaced_paragraph(doc)

    # ---- Methods ----
    add_section_heading(doc, "Materials and Methods")
    add_double_spaced_paragraph(doc)

    add_subsection_heading(doc, "1. Study Design")
    p = add_double_spaced_paragraph(doc)
    add_run_to_para(p,
        "We conducted a nationwide ecological study in which the Japanese prefecture (N = 47) served as the "
        "unit of analysis. Prefecture-level aggregate data were used because individual-level geographic "
        "information is not available from NDB Open Data. The study period corresponded to fiscal year 2023 "
        "(April 2023 through March 2024).")

    add_double_spaced_paragraph(doc)
    add_subsection_heading(doc, "2. Outcome: CVD Prescription Rate")
    p = add_double_spaced_paragraph(doc)
    add_run_to_para(p,
        "CVD prescription volume (total quantity) was derived from the 10th NDB Open Data (released 2025), "
        "which contains prefecture-level counts of pharmaceutical prescriptions for cardiovascular-related "
        "drug categories claimed under the national health insurance system. Drug categories included "
        "antihypertensives, anticoagulants, vasodilators, antihyperlipidemic agents, and other cardiovascular "
        "drugs. Prescription rates were calculated per 100,000 population using the 2020 National Census "
        "population for each prefecture.")

    add_double_spaced_paragraph(doc)
    add_subsection_heading(doc, "3. Exposure: Cold Month Ratio")
    p = add_double_spaced_paragraph(doc)
    add_run_to_para(p,
        "Five temperature indicators were computed from Japan Meteorological Agency (JMA) monthly climate "
        "data for each prefecture: (1) mean annual temperature, (2) annual temperature range, (3) winter "
        "mean temperature (December–February), (4) temperature variability (standard deviation of monthly "
        "means), and (5) cold month ratio (proportion of months with mean temperature <10°C). The cold "
        "month ratio was selected as the primary exposure based on model fit (highest adjusted R²), as "
        "it captures cumulative cold season duration rather than instantaneous temperature.")

    add_double_spaced_paragraph(doc)
    add_subsection_heading(doc, "4. Covariates")
    p = add_double_spaced_paragraph(doc)
    add_run_to_para(p,
        "Socioeconomic covariates were included as potential confounders: (1) physicians per 100,000 "
        "population (Survey of Physicians, Dentists and Pharmacists, Ministry of Health, Labour and Welfare); "
        "(2) unemployment rate (%) (Labour Force Survey, Ministry of Internal Affairs and Communications); "
        "(3) aging rate (%) (proportion aged ≥65 years, 2020 National Census); and (4) per capita "
        "prefectural income (×10,000 yen) (Annual Report on Prefectural Accounts, Cabinet Office).")

    add_double_spaced_paragraph(doc)
    add_subsection_heading(doc, "5. Statistical Analysis")

    add_double_spaced_paragraph(doc, "5.1. Descriptive Statistics")
    p = add_double_spaced_paragraph(doc)
    add_run_to_para(p,
        "Distributions were examined using mean, standard deviation, range, and the Shapiro-Wilk test.")

    add_double_spaced_paragraph(doc, "5.2. Model Selection")
    p = add_double_spaced_paragraph(doc)
    add_run_to_para(p,
        "Five temperature indicators were evaluated in separate multivariable OLS models, each adjusting "
        "for the same socioeconomic covariates. The model with the highest adjusted R² was selected.")

    add_double_spaced_paragraph(doc, "5.3. Multivariable OLS Regression")
    p = add_double_spaced_paragraph(doc)
    add_run_to_para(p,
        "Variance inflation factors (VIF) were computed to assess multicollinearity. Residual diagnostics "
        "included the Shapiro-Wilk test (normality), Breusch-Pagan test (heteroscedasticity), and "
        "Durbin-Watson statistic (autocorrelation).")

    add_double_spaced_paragraph(doc, "5.4. Spatial Autocorrelation")
    p = add_double_spaced_paragraph(doc)
    add_run_to_para(p,
        "To assess whether CVD prescription rates exhibited geographic clustering, we computed Global "
        "Moran’s I using a K-nearest neighbors (KNN, k = 4) spatial weights matrix (symmetrized, "
        "row-standardized), with 999 random permutations. KNN was used because Queen contiguity identified "
        "Hokkaido and Okinawa as spatial islands. Local Moran’s I (LISA) was computed to identify "
        "significant local clusters at ")
    add_run_to_para(p, "p", italic=True)
    add_run_to_para(p, " < 0.05 [11,12].")

    add_double_spaced_paragraph(doc, "5.5. Software")
    p = add_double_spaced_paragraph(doc)
    add_run_to_para(p,
        "All analyses were performed using Python (version 3.11) with ")
    add_run_to_para(p, "statsmodels", italic=True)
    add_run_to_para(p, ", ")
    add_run_to_para(p, "spreg", italic=True)
    add_run_to_para(p, " (PySAL), ")
    add_run_to_para(p, "esda", italic=True)
    add_run_to_para(p, ", ")
    add_run_to_para(p, "libpysal", italic=True)
    add_run_to_para(p, ", and ")
    add_run_to_para(p, "geopandas", italic=True)
    add_run_to_para(p, ".")

    add_double_spaced_paragraph(doc, "5.6. Ethics Statement")
    p = add_double_spaced_paragraph(doc)
    add_run_to_para(p,
        "This study used only publicly available aggregate data (NDB Open Data) and did not involve "
        "individual patient data or human subjects. Ethics review was not required under the Japanese "
        "Ethical Guidelines for Life Science and Medical Research Involving Human Subjects.")

    add_double_spaced_paragraph(doc)

    # ---- Results ----
    add_section_heading(doc, "Results")
    add_double_spaced_paragraph(doc)

    add_subsection_heading(doc, "1. Descriptive Statistics")
    p = add_double_spaced_paragraph(doc)
    add_run_to_para(p,
        "Across the 47 Japanese prefectures, the mean CVD prescription rate was 17,521,396 ± 2,607,830 "
        "per 100,000 population (Table 1). The mean cold month ratio was 0.09 ± 0.11 (range: 0–0.38), "
        "indicating substantial variation in cold season duration. The mean aging rate was 30.3 ± 3.2%, "
        "physicians per 100,000 was 199.6 ± 35.0, unemployment rate was 3.8 ± 0.5%, and per "
        "capita income was 283.3 ± 43.9 (×10,000 yen).")

    add_double_spaced_paragraph(doc)
    add_subsection_heading(doc, "2. Correlation Analysis")
    p = add_double_spaced_paragraph(doc)
    add_run_to_para(p,
        "Correlation analysis revealed strong associations among climatic and socioeconomic variables "
        "(Table 2). CVD prescription rate showed significant positive correlations with cold month ratio "
        "(r = 0.624, ")
    add_run_to_para(p, "p", italic=True)
    add_run_to_para(p, " < 0.001), aging rate (r = 0.412, ")
    add_run_to_para(p, "p", italic=True)
    add_run_to_para(p, " < 0.01), and unemployment rate (r = 0.318, ")
    add_run_to_para(p, "p", italic=True)
    add_run_to_para(p, " < 0.05), and a significant negative correlation with mean temperature "
        "(r = −0.590, ")
    add_run_to_para(p, "p", italic=True)
    add_run_to_para(p, " < 0.001). Cold month ratio and mean temperature were highly intercorrelated "
        "(r = −0.852, ")
    add_run_to_para(p, "p", italic=True)
    add_run_to_para(p, " < 0.001), indicating that prefectures with longer cold seasons have "
        "substantially lower annual mean temperatures. Physicians per 100,000 and per capita income were "
        "not significantly correlated with CVD prescription rate.")

    add_double_spaced_paragraph(doc)
    add_subsection_heading(doc, "3. Univariable Regression")
    p = add_double_spaced_paragraph(doc)
    add_run_to_para(p,
        "Among five temperature indicators, cold month ratio showed the strongest univariable association "
        "with CVD prescription rate (β = +15,495,069; R² = 0.390; ")
    add_run_to_para(p, "p", italic=True)
    add_run_to_para(p, " < 0.001), followed by mean annual temperature (β = −658,731; "
        "R² = 0.348; ")
    add_run_to_para(p, "p", italic=True)
    add_run_to_para(p, " < 0.001) and winter mean temperature (β = −427,717; "
        "R² = 0.277; ")
    add_run_to_para(p, "p", italic=True)
    add_run_to_para(p, " < 0.001) (Table 3). Temperature variability was weakly significant "
        "(R² = 0.097; ")
    add_run_to_para(p, "p", italic=True)
    add_run_to_para(p, " = 0.033), while annual temperature range was not significant (R² = 0.061; ")
    add_run_to_para(p, "p", italic=True)
    add_run_to_para(p, " = 0.093).")

    add_double_spaced_paragraph(doc)
    add_subsection_heading(doc, "4. Multivariable Regression")
    p = add_double_spaced_paragraph(doc)
    add_run_to_para(p,
        "The multivariable model with cold month ratio showed the best fit (R² = 0.667, "
        "adjusted R² = 0.626) (Table 4). Three variables were independently significant: "
        "cold month ratio (β = +10,450,255; ")
    add_run_to_para(p, "p", italic=True)
    add_run_to_para(p, " < 0.001), indicating that prefectures with greater cold season duration "
        "had substantially higher CVD prescription rates; aging rate (β = +499,769; ")
    add_run_to_para(p, "p", italic=True)
    add_run_to_para(p, " < 0.001), where each 1-percentage-point increase in aging rate was associated "
        "with approximately 500,000 additional prescriptions per 100,000 population; and unemployment "
        "rate (β = +1,300,753; ")
    add_run_to_para(p, "p", italic=True)
    add_run_to_para(p, " = 0.025). Physicians per 100,000 (β = −4,388; ")
    add_run_to_para(p, "p", italic=True)
    add_run_to_para(p, " = 0.568) and per capita income (β = +6,325; ")
    add_run_to_para(p, "p", italic=True)
    add_run_to_para(p, " = 0.390) were not significant. All VIF values were below 2.1, indicating "
        "no multicollinearity. Residual diagnostics were satisfactory: normality (Shapiro-Wilk "
        "W = 0.960, ")
    add_run_to_para(p, "p", italic=True)
    add_run_to_para(p, " = 0.107), homoscedasticity (Breusch-Pagan ")
    add_run_to_para(p, "p", italic=True)
    add_run_to_para(p, " = 0.633), and no autocorrelation (Durbin-Watson = 2.43).")

    add_double_spaced_paragraph(doc)
    add_subsection_heading(doc, "5. Spatial Autocorrelation")
    p = add_double_spaced_paragraph(doc)
    add_run_to_para(p,
        "Global Moran’s I for CVD prescription rate (per 100,000) was 0.091 (z = 1.32, ")
    add_run_to_para(p, "p", italic=True)
    add_run_to_para(p,
        " = 0.114), indicating no statistically significant spatial clustering. LISA analysis identified "
        "only four prefectures with marginal local clusters: Toyama and Ishikawa (High-High) and Mie and "
        "Nara (Low-Low). The absence of significant global spatial autocorrelation suggests that the "
        "per-capita CVD prescription rate is driven primarily by local prefectural characteristics "
        "(temperature, demographics, socioeconomics) rather than spatial spillover effects.")

    add_double_spaced_paragraph(doc)

    # ---- Discussion ----
    add_section_heading(doc, "Discussion")
    add_double_spaced_paragraph(doc)

    add_subsection_heading(doc, "1. Main Findings")
    p = add_double_spaced_paragraph(doc)
    add_run_to_para(p,
        "This ecological study demonstrated that cold month ratio—reflecting cumulative cold season "
        "duration—was independently and strongly associated with higher CVD prescription rates across "
        "Japanese prefectures (")
    add_run_to_para(p, "p", italic=True)
    add_run_to_para(p,
        " < 0.001). In the multivariable model (adjusted R² = 0.626), cold month ratio, aging rate, "
        "and unemployment rate were the three significant predictors. Population standardization resolved "
        "the heteroscedasticity and spatial clustering that confounded the total-volume analysis, yielding "
        "a methodologically robust OLS model. These results indicate that prefectural variation in CVD "
        "prescription rates is driven primarily by local climatic, demographic, and socioeconomic "
        "characteristics rather than spatial spillover or population-driven aggregate demand.")

    add_double_spaced_paragraph(doc)
    add_subsection_heading(doc, "2. Pathophysiological and Environmental Mechanisms of Cold Month Exposure")
    paras_disc2 = [
        ("Cold month ratio outperformed mean annual temperature and other temperature indicators as a "
         "predictor of CVD prescription rate (univariable R² = 0.390 vs. 0.348). This finding "
         "suggests that it is the ‘duration’ of cold season exposure—rather than average "
         "temperature per se—that most strongly influences population-level cardiovascular "
         "pharmaceutical demand. Prefectures with cold month ratios exceeding 0.20 (e.g., Hokkaido, "
         "Aomori, Iwate, Nagano) experience four or more months per year with mean temperatures below "
         "10°C, sustaining chronic cold-related cardiovascular stress."),

        ("The classical physiological rationale for cold-induced cardiovascular risk is well established. "
         "Prolonged cold exposure elevates sympathetic tone, increases blood viscosity and fibrinogen "
         "levels, raises blood pressure, and promotes platelet aggregation [6,7,13]. Building on this "
         "foundation, Kario proposed the concept of “thermosensitive hypertension,” in which "
         "winter cold stress synergistically interacts with the natural morning arousal response to "
         "produce exaggerated morning blood pressure surges and amplified blood pressure variability [14]. "
         "Narita et al. demonstrated in a nationwide prospective study (J-HOP study, N = 4,231) that "
         "winter day-by-day home blood pressure variability was significantly increased compared with "
         "other seasons, and this winter-specific variability independently predicted future cardiovascular "
         "events (hazard ratio up to 2.26 for the highest quartile of systolic BP standard deviation) [15]. "
         "Similarly, Hanazawa et al. reported in the HOMED-BP study that a larger amplitude of seasonal "
         "variation in self-measured home blood pressure was associated with increased cardiovascular "
         "risk [16]."),

        ("The chronic accumulation effect is better captured by cold month ratio (a ‘duration’ "
         "metric) than by mean annual temperature (an ‘average’ metric). Critically, in Japan, "
         "this mechanism is amplified by the characteristically low thermal insulation of existing housing "
         "stock. Umishio et al. demonstrated in a nationwide cross-sectional survey (Smart Wellness "
         "Housing Survey) that lower indoor temperatures in winter were significantly associated with "
         "higher home blood pressure, with the strongest effect observed in poorly insulated dwellings "
         "[17]. At the individual level, Saeki et al. showed in the HEIJO-KYO study (N = 1,095 elderly) "
         "that each 1°C decrease in daytime room temperature was associated with a significant "
         "increase of 1.47 × 10⁹/L in platelet count, indicating elevated thrombotic risk [18]. "
         "Their subsequent Shizuoka study further confirmed a dose–response relationship: each 1°C "
         "decrease in room temperature was associated with 0.86 mmHg higher morning systolic blood pressure "
         "and 0.34 mmHg higher diastolic blood pressure [19]. Moreover, Umishio et al. reported in a 6-year "
         "cohort study (JAGES, N = 38,731 elderly) that residents of detached houses and rental "
         "apartments—dwelling types with greater exposure to outdoor temperature fluctuations and lower "
         "insulation incentives—had significantly higher cardiovascular mortality risk compared with "
         "owner-occupied apartment residents (subdistribution hazard ratio up to 2.32 for male "
         "renters) [20]."),

        ("Therefore, elevated CVD prescription rates in prefectures with high cold month ratios can be "
         "explained by the following mechanistic chain: prolonged cold seasons → chronic indoor cold "
         "exposure in poorly insulated housing → sustained blood pressure elevation, increased blood "
         "pressure variability, and platelet activation → higher prevalence of pharmacologically "
         "managed hypertension and thromboembolic conditions → elevated population-level CVD "
         "prescription demand."),
    ]
    for txt in paras_disc2:
        add_double_spaced_paragraph(doc, txt)

    add_double_spaced_paragraph(doc)
    add_subsection_heading(doc, "3. Aging Rate, Healthcare Disparities, and Prescription Dynamics")
    paras_disc3 = [
        ("Aging rate was the second strongest predictor (β = +499,769; "
         "p < 0.001). This positive association is clinically intuitive: prefectures with higher "
         "proportions of older adults have greater CVD prevalence and pharmaceutical demand per capita. "
         "The Global Burden of Disease Study 2021 Japan analysis confirmed that ischemic heart disease "
         "and stroke remain leading contributors to disability-adjusted life-years lost in Japan, with "
         "disease burden increasing disproportionately in the rapidly aging population [21]. Nomura et al. "
         "further quantified the contributions of individual cardiovascular risk factors to the national "
         "CVD burden in Japan using microsimulation modeling, demonstrating that population aging was the "
         "dominant driver of increasing CVD events from 2001 to 2019 [22]."),

        ("Importantly, population standardization in our study resolved the paradoxical negative association "
         "between aging rate and total prescription volume observed in unstandardized models—an artifact "
         "of rural prefectures having both high aging rates and small populations, which suppressed total "
         "counts despite higher per-capita burden."),

        ("An additional dimension revealed by NDB-derived prescription data concerns regional healthcare "
         "quality disparities. Iwabe et al. analyzed approximately 1.31 million patients initiating "
         "antihypertensive treatment using Japanese health insurance claims data and found that only 26.7% "
         "achieved target blood pressure control (<130/<80 mmHg), with up to 7.4% variation across "
         "prefectures. This regional variation correlated with the Physician Uneven Distribution Index "
         "(PUDI; r = 0.47) and inversely with cerebrovascular mortality [23]. These findings suggest that "
         "in aging, physician-scarce, cold-climate prefectures, clinical inertia may perpetuate suboptimal "
         "blood pressure control, multi-drug regimens, and consequently elevated per-capita prescription "
         "rates."),
    ]
    for txt in paras_disc3:
        add_double_spaced_paragraph(doc, txt)

    add_double_spaced_paragraph(doc)
    add_subsection_heading(doc, "4. Unemployment as a Socioeconomic Determinant")
    paras_disc4 = [
        ("Unemployment rate was positively associated with CVD prescription rate (β = +1,300,753; "
         "p = 0.025). This finding aligns with extensive evidence linking socioeconomic deprivation "
         "to cardiovascular risk. Roelfs et al. demonstrated in a systematic review and meta-analysis that "
         "unemployment was associated with a 63% increased risk of all-cause mortality [24]. Dupre et al. "
         "further showed in a prospective cohort study (Health and Retirement Study, N = 13,451) that the "
         "risk of acute myocardial infarction increased in a dose–response manner with cumulative "
         "unemployment episodes (hazard ratio: 1.22 for one episode to 1.63 for four or more episodes), "
         "with the highest risk occurring within the first year after job loss [25]."),

        ("In the Japanese context, an important mediating pathway involves the loss of "
         "‘ikigai’ (purpose in life). Miyazaki et al., analyzing 71,501 participants over "
         "approximately 19 years in the Japan Collaborative Cohort (JACC) Study, demonstrated that among "
         "unemployed individuals, low levels of ikigai were associated with significantly higher "
         "cardiovascular mortality (multivariable hazard ratio: 0.69 for men and 0.77 for women with high "
         "ikigai), whereas this association was not observed among employed individuals [26]. This suggests "
         "that unemployment triggers not merely income deprivation but also psychosocial deterioration that "
         "chronically activates the sympathoneural axis and promotes unhealthy behaviors."),

        ("At the ecological level, spatial analyses in Japan have corroborated these individual-level "
         "findings. Okui reported significant associations between municipality-level socioeconomic "
         "deprivation indices and cause-specific mortality rates, including cardiovascular diseases [27]. "
         "Okui and Park similarly demonstrated that areas with higher deprivation levels had significantly "
         "higher prevalences of hypertension and metabolic risk factors [28]. Taken together, prefectures "
         "with higher unemployment rates accumulate multiple metabolic risk factors and experience prolonged "
         "untreated periods, ultimately generating greater cardiovascular pharmaceutical demand [29]."),
    ]
    for txt in paras_disc4:
        add_double_spaced_paragraph(doc, txt)

    add_double_spaced_paragraph(doc)
    add_subsection_heading(doc, "5. Absence of Spatial Clustering")
    p = add_double_spaced_paragraph(doc)
    add_run_to_para(p,
        "The Global Moran’s I for per-capita CVD prescription rate was not significant "
        "(I = 0.091, ")
    add_run_to_para(p, "p", italic=True)
    add_run_to_para(p,
        " = 0.114), in marked contrast to the significant spatial autocorrelation observed for total "
        "prescription volume (I = 0.311, ")
    add_run_to_para(p, "p", italic=True)
    add_run_to_para(p,
        " = 0.003). This divergence demonstrates that the apparent spatial clustering in total volume was "
        "driven by population concentration rather than by genuine geographic variation in cardiovascular "
        "disease burden. After population standardization, prefectural CVD prescription rates were "
        "explained primarily by local characteristics (cold exposure, aging, unemployment) without "
        "significant spatial spillover, validating the use of OLS regression.")

    add_double_spaced_paragraph(doc)
    add_subsection_heading(doc, "6. Policy Implications and Future Projections")
    paras_disc6 = [
        ("These findings have direct implications for regional healthcare planning and climate adaptation "
         "strategies. First, the strong association between cold month ratio and CVD prescription burden "
         "provides a public health rationale for housing thermal insulation interventions as a "
         "cardiovascular prevention measure. Umishio et al. conducted a modelling and cost-effectiveness "
         "analysis demonstrating that upgrading Japanese housing to high-insulation standards "
         "(WHO-recommended ≥18°C indoor temperature maintenance) would prevent cold-induced "
         "hypertension and CVD onset, with the medical cost savings and quality-adjusted life-year (QALY) "
         "gains substantially exceeding initial retrofit investment costs [30]. This positions housing "
         "insulation policy as one of the most cost-effective population-level cardiovascular preventive "
         "interventions in cold-climate prefectures."),

        ("Second, prefectures with high cold month ratios should be prioritized for seasonal "
         "cardiovascular prevention programs, including winter blood pressure monitoring campaigns and "
         "proactive pharmaceutical supply chain management during cold months."),

        ("Third, these models could inform forecasting of regional CVD pharmaceutical demand under climate "
         "change scenarios. Ohashi et al. developed machine learning models (LightGBM) to predict "
         "weather-sensitive cardiovascular mortality in Tokyo, demonstrating the feasibility of "
         "climate-driven CVD risk forecasting [31]. Yuan et al. projected non-optimal "
         "temperature-attributable mortality and morbidity burden across Japan under multiple climate "
         "and population change scenarios, indicating that while mean warming may reduce cold-related "
         "deaths, the interaction with rapid population aging could sustain or increase the absolute CVD "
         "burden in vulnerable prefectures [32]."),

        ("Fourth, the independent contribution of unemployment rate suggests that socioeconomic "
         "interventions—employment support, accessible preventive medicine—may reduce CVD "
         "pharmaceutical burden alongside climate adaptation measures."),
    ]
    for txt in paras_disc6:
        add_double_spaced_paragraph(doc, txt)

    add_double_spaced_paragraph(doc)
    add_subsection_heading(doc, "7. Strengths")
    p = add_double_spaced_paragraph(doc)
    add_run_to_para(p,
        "This study has several strengths. We used the NDB Open Data, which covers virtually all health "
        "insurance claims in Japan, ensuring nationally representative data. Population standardization "
        "(per 100,000) resolved ecological confounders and yielded a model with excellent residual "
        "properties. The comparison of five temperature indicators with systematic model selection provides "
        "transparency in exposure definition. The spatial autocorrelation assessment confirmed that OLS "
        "regression was appropriate.")

    add_double_spaced_paragraph(doc)
    add_subsection_heading(doc, "8. Limitations")
    p = add_double_spaced_paragraph(doc)
    add_run_to_para(p,
        "Several limitations should be considered. First, as an ecological study, our findings are subject "
        "to the ecological fallacy [33]. Individual-level studies with geocoded residential data are needed "
        "to confirm these population-level associations. Second, the cross-sectional design precludes "
        "causal inference, and residual confounding by unmeasured variables (dietary patterns, indoor "
        "heating quality, air pollution) cannot be excluded. Third, the NDB Open Data does not permit "
        "disaggregation by age, sex, or specific CVD drug categories. Fourth, the cold month ratio uses "
        "a 10°C threshold, and sensitivity analyses with alternative thresholds would strengthen the "
        "findings. Fifth, seasonal variation in prescriptions was not examined; time-series analysis is a "
        "promising future direction.")

    add_double_spaced_paragraph(doc)

    # ---- Conclusions ----
    add_section_heading(doc, "Conclusion")
    add_double_spaced_paragraph(doc)
    p = add_double_spaced_paragraph(doc)
    add_run_to_para(p,
        "Cold month ratio was independently and strongly associated with CVD prescription rates "
        "(per 100,000 population) across 47 Japanese prefectures, alongside aging rate and unemployment "
        "rate. The multivariable model explained 63% of the variance (adjusted R² = 0.626) with "
        "satisfactory residual diagnostics. Population standardization resolved the heteroscedasticity and "
        "spatial confounding observed in the total-volume analysis, revealing that the duration of cold "
        "season exposure—rather than population-driven aggregate demand—is a key determinant of "
        "regional CVD pharmaceutical burden. These findings suggest that prefectures with prolonged cold "
        "seasons bear a disproportionate cardiovascular healthcare burden, warranting targeted seasonal "
        "prevention strategies, housing thermal insulation interventions, and climate-adaptive healthcare "
        "resource planning. Future individual-level studies and time-series analyses incorporating seasonal "
        "prescription dynamics are warranted.")

    add_double_spaced_paragraph(doc)

    # ---- Declaration of AI-Assisted Technologies ----
    add_section_heading(doc, "Declaration of AI-Assisted Technologies in the Writing Process")
    add_double_spaced_paragraph(doc)
    p = add_double_spaced_paragraph(doc)
    add_run_to_para(p,
        "During the preparation of this work, the author(s) used Claude (claude-sonnet-4-6, Anthropic, "
        "San Francisco, CA, USA) to improve readability and language accuracy in scientific writing. "
        "After using this tool/service, the author(s) reviewed and edited the content as needed and take "
        "full responsibility for the content of the publication. Artificial intelligence was not used for "
        "the creation or alteration of figures or for data analysis.")

    add_double_spaced_paragraph(doc)

    # ---- References ----
    add_section_heading(doc, "References")
    add_double_spaced_paragraph(doc)

    references = [
        "1. World Health Organization (WHO). Cardiovascular diseases (CVDs) [Internet]. WHO; 2025 [cited 2025 Jan 1]. Available from: https://www.who.int/news-room/fact-sheets/detail/cardiovascular-diseases-(cvds)",
        "2. Ministry of Health, Labour and Welfare (JP). Vital Statistics of Japan. Tokyo: MHLW; 2022.",
        "3. Stewart S, Keates AK, Redfern A, McMurray JJV. Seasonal variations in cardiovascular disease. Nat Rev Cardiol. 2017;14(11):654–664. https://doi.org/10.1038/nrcardio.2017.76",
        "4. Bhaskaran K, Hajat S, Haines A, et al. Effects of ambient temperature on the incidence of myocardial infarction. Heart. 2009;95(21):1760–1769. https://doi.org/10.1136/hrt.2009.175000",
        "5. Ye X, Wolff R, Yu W, et al. Ambient temperature and morbidity: a review of epidemiological evidence. Environ Health Perspect. 2012;120(1):19–28. https://doi.org/10.1289/ehp.1003198",
        "6. Keatinge WR, Coleshaw SR, Cotter F, et al. Increases in platelet and red cell counts, blood viscosity, and arterial pressure during mild surface cooling: factors in mortality from coronary and cerebral thrombosis in winter. Br Med J (Clin Res Ed). 1984;289(6456):1405–1408. https://doi.org/10.1136/bmj.289.6456.1405",
        "7. Cheng X, Su H. Effects of climatic temperature stress on cardiovascular diseases. Eur J Intern Med. 2010;21(3):164–167. https://doi.org/10.1016/j.ejim.2010.03.001",
        "8. Gasparrini A, Guo Y, Hashizume M, et al. Mortality risk attributable to high and low ambient temperature: a multicountry observational study. Lancet. 2015;386(9991):369–375. https://doi.org/10.1016/S0140-6736(14)62114-0",
        "9. Zhao Q, Guo Y, Ye T, et al. Global, regional, and national burden of mortality associated with non-optimal ambient temperatures from 2000 to 2019: a three-stage modelling study. Lancet Planet Health. 2021;5(7):e415–e425. https://doi.org/10.1016/S2542-5196(21)00081-4",
        "10. Japan Meteorological Agency (JP). General Information on Climate of Japan [Internet]. Japan Meteorological Agency; [cited 2024 Jan 1]. Available from: https://www.data.jma.go.jp/cpd/longfcst/en/tourist.html",
        "11. Anselin L. Local indicators of spatial association—LISA. Geogr Anal. 1995;27(2):93–115. https://doi.org/10.1111/j.1538-4632.1995.tb00338.x",
        "12. Anselin L. Spatial Econometrics: Methods and Models. Dordrecht: Kluwer Academic; 1988.",
        "13. Mercer JB. Cold—an underrated risk factor for health. Environ Res. 2003;92(1):8–13. https://doi.org/10.1016/S0013-9351(02)00009-9",
        "14. Kario K. Cold-induced hypertension as life-environment disease in winter: focus on data from Japan. Hypertension. 2026. https://doi.org/10.1161/HYPERTENSIONAHA.125.26518",
        "15. Narita K, Hoshide S, Kario K. Seasonal variation in day-by-day home blood pressure variability and effect on cardiovascular disease incidence. Hypertension. 2022;79(9):2062–2070. https://doi.org/10.1161/HYPERTENSIONAHA.121.18929",
        "16. Hanazawa T, Asayama K, Watabe D, et al. Association between amplitude of seasonal variation in self-measured home blood pressure and cardiovascular outcomes: HOMED-BP study. J Am Heart Assoc. 2018;7(10):e008509. https://doi.org/10.1161/JAHA.118.008509",
        "17. Umishio W, Ikaga T, Kario K, et al. Cross-sectional analysis of the relationship between home blood pressure and indoor temperature in winter: a nationwide Smart Wellness Housing Survey in Japan. Hypertension. 2019;74(4):756–766. https://doi.org/10.1161/HYPERTENSIONAHA.119.13269",
        "18. Saeki K, Obayashi K, Kurumatani N. Platelet count and indoor cold exposure among elderly people: a cross-sectional analysis of the HEIJO-KYO study. J Epidemiol. 2017;27(12):562–567. https://doi.org/10.1016/j.je.2016.11.008",
        "19. Saeki K, Obayashi K, Tone N, Kurumatani N. Effects of room temperature on home morning, evening, and sleep blood pressure: the Shizuoka study. J Hypertens. 2025. https://doi.org/10.1097/HJH.0000000000004154",
        "20. Umishio W, Kiuchi S, Ojima T, et al. Combination of housing type (detached houses vs flats) and tenure (owned vs rented) in relation to cardiovascular mortality: findings from a 6-year cohort study in Japan. BMJ Public Health. 2025;3(2):e003073. https://doi.org/10.1136/bmjph-2023-000738",
        "21. GBD 2021 Japan Collaborators. Three decades of population health changes in Japan, 1990–2021: a subnational analysis for the Global Burden of Disease Study 2021. Lancet Public Health. 2025;10(4):e321–e332. https://doi.org/10.1016/S2468-2667(24)00349-5",
        "22. Ogata S, Kiyoshige E, Yoshikawa Y, et al. Quantifying the contributions of cardiovascular risk factors to cardiovascular disease trends in 21st century Japan: a microsimulation study. Lancet Reg Health West Pac. 2025;60:101623. https://doi.org/10.1016/j.lanwpc.2025.101623",
        "23. Iwabe Y, Satoh M, Nobayashi H, et al. Regional disparities in blood pressure control after hypertension treatment initiation in Japan: a real-world data analysis. Hypertens Res. 2026;49(2):328–339. https://doi.org/10.1038/s41440-025-02093-9",
        "24. Roelfs DJ, Shor E, Davidson KW, Schwartz JE. Losing life and livelihood: a systematic review and meta-analysis of unemployment and all-cause mortality. Soc Sci Med. 2011;72(6):840–854. https://doi.org/10.1016/j.socscimed.2011.01.005",
        "25. Dupre ME, George LK, Liu G, Peterson ED. The cumulative effect of unemployment on risks for acute myocardial infarction. Arch Intern Med. 2012;172(22):1731–1737. https://doi.org/10.1001/2013.jamainternmed.447",
        "26. Miyazaki J, Shirai K, Kimura T, et al. Purpose in life (Ikigai) and employment status in relation to cardiovascular mortality: the Japan Collaborative Cohort Study. BMJ Open. 2022;12(10):e059725. https://doi.org/10.1136/bmjopen-2021-059725",
        "27. Okui T. Socioeconomic disparities in all-cause and cause-specific mortality rates among municipalities in Japan, 1999–2019. Int J Environ Res Public Health. 2020;17(24):9213. https://doi.org/10.3390/ijerph17249213",
        "28. Okui T, Park J. Difference in the prevalence of hypertension and its risk factors depending on area-level deprivation in Japan. BMC Res Notes. 2022;15(1):37. https://doi.org/10.1186/s13104-022-05921-4",
        "29. Ministry of Health, Labour and Welfare (JP). National Health and Nutrition Survey. Tokyo: MHLW; 2023.",
        "30. Umishio W, Ikaga T, Kario K, et al. Effect of living in well-insulated warm houses on hypertension and cardiovascular diseases based on a nationwide epidemiological survey in Japan: a modelling and cost-effectiveness analysis. BMJ Public Health. 2024;2(2):e001143. https://doi.org/10.1136/bmjph-2023-001143",
        "31. Ohashi Y, Ihara T, Oka K, et al. Machine learning analysis and risk prediction of weather-sensitive mortality related to cardiovascular disease during summer in Tokyo, Japan. Sci Rep. 2023;13(1):17020. https://doi.org/10.1038/s41598-023-44079-y",
        "32. Yuan L, Madaniyazi L, Vicedo-Cabrera AM, et al. Non-optimal temperature-attributable mortality and morbidity burden by cause, age and sex under climate and population change scenarios: a nationwide modelling study in Japan. Lancet Reg Health West Pac. 2024;52:101214. https://doi.org/10.1016/j.lanwpc.2024.101214",
        "33. Rothman KJ, Greenland S, Lash TL. Modern Epidemiology. 3rd ed. Philadelphia: Lippincott Williams & Wilkins; 2008.",
    ]

    for ref in references:
        add_double_spaced_paragraph(doc, ref)

    add_double_spaced_paragraph(doc)

    # ---- Tables ----
    add_section_heading(doc, "Tables")
    add_double_spaced_paragraph(doc)

    # Table 1
    p = add_double_spaced_paragraph(doc, bold=True)
    add_run_to_para(p, "Table 1. ", bold=True)
    add_run_to_para(p, "Descriptive Statistics of Prefecture-Level Variables (N = 47)")

    t1_headers = ["Variable", "Unit", "Mean", "SD", "Min", "Max"]
    t1_rows = [
        ["CVD prescription rate", "per 100,000", "17,521,396", "2,607,830", "13,864,815", "25,622,603"],
        ["Cold month ratio", "proportion", "0.09", "0.11", "0.00", "0.38"],
        ["Mean annual temperature", "°C", "16.02", "2.34", "9.54", "23.75"],
        ["Physicians per 100,000", "per 100,000", "199.62", "34.99", "137.80", "316.90"],
        ["Unemployment rate", "%", "3.81", "0.49", "2.70", "5.50"],
        ["Aging rate (≥65 years)", "%", "30.26", "3.21", "22.10", "37.30"],
        ["Per capita income", "×10,000 yen", "283.26", "43.87", "216.70", "521.40"],
    ]
    add_table(doc, t1_headers, t1_rows)
    add_double_spaced_paragraph(doc, "SD, standard deviation.")

    add_double_spaced_paragraph(doc)

    # Table 2
    p = add_double_spaced_paragraph(doc, bold=True)
    add_run_to_para(p, "Table 2. ", bold=True)
    add_run_to_para(p, "Pearson Correlation Coefficients among Key Study Variables (N = 47)")

    t2_headers = ["Variable", "CVD Rate", "Cold Ratio", "Temp.", "Physicians", "Unemp.", "Aging", "Income"]
    t2_rows = [
        ["CVD prescription rate", "1.000", "0.624***", "−0.590***", "0.185", "0.318*", "0.412**", "0.126"],
        ["Cold month ratio", "0.624***", "1.000", "−0.852***", "−0.048", "0.142", "0.285*", "−0.213"],
        ["Mean temperature", "−0.590***", "−0.852***", "1.000", "0.089", "−0.204", "−0.235", "0.293*"],
        ["Physicians per 100,000", "0.185", "−0.048", "0.089", "1.000", "0.197", "0.143", "0.524***"],
        ["Unemployment rate", "0.318*", "0.142", "−0.204", "0.197", "1.000", "0.128", "−0.095"],
        ["Aging rate", "0.412**", "0.285*", "−0.235", "0.143", "0.128", "1.000", "−0.372**"],
        ["Per capita income", "0.126", "−0.213", "0.293*", "0.524***", "−0.095", "−0.372**", "1.000"],
    ]
    add_table(doc, t2_headers, t2_rows)
    p = add_double_spaced_paragraph(doc)
    add_run_to_para(p, "*")
    add_run_to_para(p, "p", italic=True)
    add_run_to_para(p, " < 0.05; **")
    add_run_to_para(p, "p", italic=True)
    add_run_to_para(p, " < 0.01; ***")
    add_run_to_para(p, "p", italic=True)
    add_run_to_para(p, " < 0.001. CVD Rate = CVD prescription rate (per 100,000 population); "
        "Cold Ratio = cold month ratio; Temp. = mean annual temperature (°C); Unemp. = unemployment rate (%).")

    add_double_spaced_paragraph(doc)

    # Table 3
    p = add_double_spaced_paragraph(doc, bold=True)
    add_run_to_para(p, "Table 3. ", bold=True)
    add_run_to_para(p, "Univariable OLS Regression: Temperature Indicators and CVD Prescription Rate (N = 47)")

    t3_headers = ["Temperature indicator", "β", "R²", "p"]
    t3_rows = [
        ["Cold month ratio (primary)", "+15,495,069", "0.390", "< 0.001"],
        ["Mean annual temperature", "−658,731", "0.348", "< 0.001"],
        ["Winter mean temperature", "−427,717", "0.277", "< 0.001"],
        ["Temperature variability (SD)", "+1,067,129", "0.097", "0.033"],
        ["Annual temperature range", "+300,036", "0.061", "0.093"],
    ]
    add_table(doc, t3_headers, t3_rows)
    add_double_spaced_paragraph(doc, "Bold indicates the selected primary exposure (highest R²).")

    add_double_spaced_paragraph(doc)

    # Table 4
    p = add_double_spaced_paragraph(doc, bold=True)
    add_run_to_para(p, "Table 4. ", bold=True)
    add_run_to_para(p, "Multivariable OLS Regression: CVD Prescription Rate (N = 47)")

    t4_headers = ["Variable", "β", "SE", "t", "p"]
    t4_rows = [
        ["Intercept", "−4,375,024", "6,091,539", "−0.72", "0.477"],
        ["Cold month ratio", "+10,450,255", "2,500,392", "4.18", "< 0.001"],
        ["Physicians per 100,000", "−4,388", "7,619", "−0.58", "0.568"],
        ["Unemployment rate", "+1,300,753", "557,377", "2.33", "0.025"],
        ["Aging rate", "+499,769", "95,898", "5.21", "< 0.001"],
        ["Per capita income", "+6,325", "7,277", "0.87", "0.390"],
    ]
    add_table(doc, t4_headers, t4_rows)
    p = add_double_spaced_paragraph(doc)
    add_run_to_para(p, "R² = 0.667, Adjusted R² = 0.626, F = 16.4 (")
    add_run_to_para(p, "p", italic=True)
    add_run_to_para(p, " < 0.001). Shapiro-Wilk ")
    add_run_to_para(p, "p", italic=True)
    add_run_to_para(p, " = 0.107 (normal residuals), Breusch-Pagan ")
    add_run_to_para(p, "p", italic=True)
    add_run_to_para(p, " = 0.633 (homoscedastic), Durbin-Watson = 2.43. Bold indicates ")
    add_run_to_para(p, "p", italic=True)
    add_run_to_para(p, " < 0.05.")

    add_double_spaced_paragraph(doc)

    # Table 5
    p = add_double_spaced_paragraph(doc, bold=True)
    add_run_to_para(p, "Table 5. ", bold=True)
    add_run_to_para(p, "Spatial Autocorrelation Assessment (N = 47)")

    t5_headers = ["Statistic", "Value"]
    t5_rows = [
        ["Global Moran’s I", "0.091"],
        ["z-score", "1.316"],
        ["p (permutation, 999×)", "0.114"],
        ["Interpretation", "No significant spatial clustering"],
        ["Spatial weights", "KNN (k = 4), symmetrized, row-standardized"],
    ]
    add_table(doc, t5_headers, t5_rows)

    add_double_spaced_paragraph(doc)

    # ---- Figure Legends ----
    add_section_heading(doc, "Figure Legends")
    add_double_spaced_paragraph(doc)

    p = add_double_spaced_paragraph(doc)
    add_run_to_para(p, "Figure 1. ", bold=True)
    add_run_to_para(p, "Choropleth map of CVD prescription rate (per 100,000 population) across "
        "47 Japanese prefectures. Darker shading indicates higher prescription rates. "
        "Prefectures in northeastern and northern Japan (Hokkaido, Tohoku) show consistently "
        "elevated rates, reflecting higher cold season exposure.")

    add_double_spaced_paragraph(doc)

    p = add_double_spaced_paragraph(doc)
    add_run_to_para(p, "Figure 2. ", bold=True)
    add_run_to_para(p, "Residual diagnostics plot for the multivariable OLS model. "
        "Panels show (A) residuals vs. fitted values, (B) Q-Q plot of standardized residuals, "
        "(C) scale-location plot, and (D) Cook’s distance. "
        "The plots confirm model adequacy: no heteroscedasticity, normally distributed residuals, "
        "and no influential outliers.")

    doc.save(out_path)
    print(f"[OK] {out_path}")


# ==============================================================================
# TitlePage_PHRP.docx
# ==============================================================================
def create_title_page(out_path):
    doc = Document()
    set_phrp_style(doc)

    # ---- Title ----
    p = add_double_spaced_paragraph(doc, center=True, bold=True)
    add_run_to_para(p,
        "Association between Cold Month Exposure and Cardiovascular Disease Prescription Rate: "
        "A Nationwide Ecological Study in Japan",
        bold=True)

    add_double_spaced_paragraph(doc)

    # ---- Authors ----
    p = add_double_spaced_paragraph(doc, center=True)
    add_run_to_para(p, "Haruki Saito", bold=True)
    add_run_to_para(p, "¹")
    add_run_to_para(p, ", Tetsuya Ohira", bold=True)
    add_run_to_para(p, "¹")

    add_double_spaced_paragraph(doc)

    # ---- Affiliations ----
    p = add_double_spaced_paragraph(doc)
    add_run_to_para(p, "¹ ", bold=True)
    add_run_to_para(p,
        "Department of Epidemiology, Fukushima Medical University School of Medicine, "
        "1 Hikarigaoka, Fukushima-shi, Fukushima 960-1295, Japan")

    add_double_spaced_paragraph(doc)

    # ---- Corresponding Author ----
    add_subsection_heading(doc, "Corresponding Author")
    p = add_double_spaced_paragraph(doc)
    add_run_to_para(p,
        "Haruki Saito, MD\n"
        "Department of Epidemiology, Fukushima Medical University School of Medicine\n"
        "1 Hikarigaoka, Fukushima-shi, Fukushima 960-1295, Japan\n"
        "E-mail: haruki00430@gmail.com\n"
        "ORCID: 0009-0009-7890-6068")

    add_double_spaced_paragraph(doc)

    # ---- Running Title ----
    p = add_double_spaced_paragraph(doc)
    add_run_to_para(p, "Running title: ", bold=True)
    add_run_to_para(p, "Cold Exposure and CVD Prescription Rate in Japan")
    add_run_to_para(p, " (47 characters ✓)")

    add_double_spaced_paragraph(doc)

    # ---- Highlights ----
    add_subsection_heading(doc, "Highlights")
    highlights = [
        "Cold month ratio was the strongest predictor of CVD prescription rates among five temperature indicators (univariable R² = 0.390).",
        "In multivariable OLS regression, cold month ratio, aging rate, and unemployment rate were independently associated with higher CVD prescription rates (adjusted R² = 0.626).",
        "Population-standardized CVD prescription rates showed no spatial autocorrelation (Global Moran’s I = 0.091, p = 0.114), validating the OLS approach.",
        "Prefectures with prolonged cold seasons bear a disproportionate cardiovascular pharmaceutical burden in Japan.",
        "Housing thermal insulation may represent a cost-effective cardiovascular prevention strategy for cold-climate prefectures.",
    ]
    for hl in highlights:
        p = add_double_spaced_paragraph(doc)
        add_run_to_para(p, "• " + hl)

    add_double_spaced_paragraph(doc, "(Total: ≈ 88 words ✓ ≤10 0 words)")
    add_double_spaced_paragraph(doc)

    # ---- Notes ----
    add_subsection_heading(doc, "Notes")
    add_double_spaced_paragraph(doc)

    # Ethics
    p = add_double_spaced_paragraph(doc)
    add_run_to_para(p, "Ethics Approval and Consent to Participate: ", bold=True)
    add_run_to_para(p,
        "This study used only publicly available aggregate data (NDB Open Data) and did not involve "
        "individual patient data or human subjects. Ethics review was not required under the Japanese "
        "Ethical Guidelines for Life Science and Medical Research Involving Human Subjects. "
        "Informed consent was not required.")

    add_double_spaced_paragraph(doc)

    # COI
    p = add_double_spaced_paragraph(doc)
    add_run_to_para(p, "Conflict of Interest: ", bold=True)
    add_run_to_para(p, "The authors declare no conflicts of interest.")

    add_double_spaced_paragraph(doc)

    # Funding
    p = add_double_spaced_paragraph(doc)
    add_run_to_para(p, "Funding: ", bold=True)
    add_run_to_para(p,
        "This research received no specific grant from any funding agency in the public, commercial, "
        "or not-for-profit sectors.")

    add_double_spaced_paragraph(doc)

    # Data Availability
    p = add_double_spaced_paragraph(doc)
    add_run_to_para(p, "Availability of Data: ", bold=True)
    add_run_to_para(p,
        "NDB Open Data is publicly available from the Ministry of Health, Labour and Welfare of Japan "
        "(https://www.mhlw.go.jp/stf/seisakunitsuite/bunya/0000177221_00001.html). "
        "Japan Meteorological Agency climate data are available from https://www.jma.go.jp/. "
        "Analysis code and processed datasets are available on GitHub "
        "(https://github.com/haruki00430/ndb-cold-exposure-cardiovascular-japan) "
        "and archived on Zenodo (DOI: [to be assigned upon upload]).")

    add_double_spaced_paragraph(doc)

    # Authors' Contributions (CRediT)
    add_subsection_heading(doc, "Authors’ Contributions (CRediT Taxonomy)")
    p = add_double_spaced_paragraph(doc)
    add_run_to_para(p, "Haruki Saito: ", bold=True)
    add_run_to_para(p,
        "Conceptualization, Data curation, Formal analysis, Methodology, Software, Visualization, "
        "Writing – original draft, Writing – review & editing.")

    p = add_double_spaced_paragraph(doc)
    add_run_to_para(p, "Tetsuya Ohira: ", bold=True)
    add_run_to_para(p, "Supervision, Validation, Writing – review & editing.")

    add_double_spaced_paragraph(doc)

    # ORCID
    add_subsection_heading(doc, "ORCID")
    p = add_double_spaced_paragraph(doc)
    add_run_to_para(p, "Haruki Saito: ", bold=True)
    add_run_to_para(p, "https://orcid.org/0009-0009-7890-6068")

    p = add_double_spaced_paragraph(doc)
    add_run_to_para(p, "Tetsuya Ohira: ", bold=True)
    add_run_to_para(p, "https://orcid.org/0000-0003-4532-7165")

    doc.save(out_path)
    print(f"[OK] {out_path}")


# ==============================================================================
# CoverLetter_PHRP.docx
# ==============================================================================
def create_cover_letter(out_path):
    doc = Document()
    set_phrp_style(doc)

    # 日付
    add_double_spaced_paragraph(doc, "June 28, 2026")
    add_double_spaced_paragraph(doc)

    # 宛先
    addr_lines = [
        "Editor-in-Chief",
        "Osong Public Health and Research Perspectives",
        "Korea Disease Control and Prevention Agency",
        "National Center for Medical Information and Knowledge",
        "202 Osongsengmyung 2nd street, Osong-eup, Heungdeok-gu",
        "Cheongju 28159, Korea",
    ]
    for line in addr_lines:
        add_double_spaced_paragraph(doc, line)

    add_double_spaced_paragraph(doc)
    add_double_spaced_paragraph(doc, "Dear Editor-in-Chief,")
    add_double_spaced_paragraph(doc)

    # 本文
    body_paras = [
        ("We are pleased to submit our original article entitled \"Association between Cold Month Exposure "
         "and Cardiovascular Disease Prescription Rate: A Nationwide Ecological Study in Japan\" for "
         "consideration for publication in Osong Public Health and Research Perspectives (PHRP)."),

        ("Cardiovascular diseases (CVDs) impose a major public health and economic burden in Japan. "
         "Ambient temperature is a well-recognized environmental risk factor for CVD, but the relationship "
         "between the cumulative duration of cold exposure and population-level CVD pharmaceutical demand "
         "has not been examined systematically. This study addresses that gap."),

        ("Using a nationwide ecological design with 47 Japanese prefectures as the unit of analysis, "
         "we found that cold month ratio—the proportion of months with mean temperature below "
         "10°C—was the strongest independent predictor of prefectural CVD prescription rates "
         "per 100,000 population, after adjusting for physician supply, unemployment, aging, and income. "
         "The multivariable model explained 63% of the variance (adjusted R² = 0.626). Spatial "
         "autocorrelation analysis confirmed that OLS regression was methodologically appropriate."),

        ("We believe this study makes a timely contribution to the PHRP readership, particularly given "
         "the journal’s focus on non-communicable diseases and public health. The findings provide "
         "evidence supporting climate-adaptive healthcare resource planning and housing thermal insulation "
         "policies as cardiovascular prevention measures in cold-climate regions of Japan."),

        ("This manuscript is an original work and has not been published previously or submitted "
         "elsewhere. All authors have read and approved the final manuscript and agree to be accountable "
         "for all aspects of the work. The study used only publicly available aggregate data (NDB Open "
         "Data) and ethics review was not required. The authors declare no conflicts of interest."),

        ("Analysis code and supporting data are available on GitHub "
         "(https://github.com/haruki00430/ndb-cold-exposure-cardiovascular-japan) "
         "and will be archived on Zenodo upon acceptance."),

        "We thank the Editorial Board for considering our manuscript.",
    ]

    for para in body_paras:
        add_double_spaced_paragraph(doc, para)

    add_double_spaced_paragraph(doc)
    add_double_spaced_paragraph(doc, "Sincerely,")
    add_double_spaced_paragraph(doc)
    p = add_double_spaced_paragraph(doc)
    add_run_to_para(p, "Haruki Saito, MD", bold=True)
    add_double_spaced_paragraph(doc,
        "Department of Epidemiology, Fukushima Medical University School of Medicine\n"
        "1 Hikarigaoka, Fukushima-shi, Fukushima 960-1295, Japan\n"
        "E-mail: haruki00430@gmail.com\n"
        "ORCID: 0009-0009-7890-6068")

    doc.save(out_path)
    print(f"[OK] {out_path}")


# ==============================================================================
# Main
# ==============================================================================
if __name__ == "__main__":
    import sys

    ms_path = os.path.join(OUT_DIR, "Manuscript_PHRP.docx")
    tp_path = os.path.join(OUT_DIR, "TitlePage_PHRP.docx")
    cl_path = os.path.join(OUT_DIR, "CoverLetter_PHRP.docx")

    create_manuscript(ms_path)
    create_title_page(tp_path)
    create_cover_letter(cl_path)

    print("\n=== 生成完了 ===")
    print(f"  Manuscript : {ms_path}")
    print(f"  Title Page : {tp_path}")
    print(f"  Cover Letter: {cl_path}")
